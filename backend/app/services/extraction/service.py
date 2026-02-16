"""
Extract text from PDF and DOCX. Uses pdfplumber for PDF, python-docx for DOCX.
Handles large files safely by limiting pages/chars.
"""
import logging
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

# Safety limits
MAX_PDF_PAGES = 50
MAX_EXTRACTED_CHARS = 100_000


class ExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


class ExtractionService:
    """Extract raw text from PDF and DOCX files."""

    @staticmethod
    def extract_from_path(file_path: str | Path) -> str:
        """
        Extract text from file. Supports .pdf and .docx.
        Raises ExtractionError on failure or unsupported type.
        """
        path = Path(file_path)
        if not path.exists():
            raise ExtractionError(f"File not found: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return ExtractionService._extract_pdf(path)
        if suffix == ".docx":
            return ExtractionService._extract_docx(path)
        raise ExtractionError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            parts: list[str] = []
            total_chars = 0
            with pdfplumber.open(path) as pdf:
                pages = min(len(pdf.pages), MAX_PDF_PAGES)
                for i in range(pages):
                    if total_chars >= MAX_EXTRACTED_CHARS:
                        break
                    page = pdf.pages[i]
                    text = page.extract_text()
                    if text:
                        parts.append(text)
                        total_chars += len(text)
            return "\n".join(parts) if parts else ""
        except Exception as e:
            logger.exception("PDF extraction failed: %s", path)
            raise ExtractionError(f"PDF extraction failed: {e}") from e

    @staticmethod
    def _extract_docx(path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(path)
            parts: list[str] = []
            total_chars = 0
            for para in doc.paragraphs:
                if total_chars >= MAX_EXTRACTED_CHARS:
                    break
                if para.text:
                    parts.append(para.text)
                    total_chars += len(para.text)
            for table in doc.tables:
                if total_chars >= MAX_EXTRACTED_CHARS:
                    break
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
                            total_chars += len(cell.text)
            return "\n".join(parts) if parts else ""
        except Exception as e:
            logger.exception("DOCX extraction failed: %s", path)
            raise ExtractionError(f"DOCX extraction failed: {e}") from e
